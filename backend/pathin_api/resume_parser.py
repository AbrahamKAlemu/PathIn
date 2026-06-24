from __future__ import annotations

import hashlib
import io
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader
from werkzeug.datastructures import FileStorage

from .errors import ApiError
from .taxonomy import CONCEPT_ALIASES, OCCUPATIONAL_TAXONOMY, normalize_title

MAX_UPLOAD_BYTES = 5 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 250_000
MAX_DOCX_UNCOMPRESSED_BYTES = 25 * 1024 * 1024
MAX_DOCX_ENTRIES = 2_000
MAX_PDF_PAGES = 200

SUPPORTED_TYPES = {
    ".pdf": {
        "mimeTypes": {"application/pdf", "application/octet-stream"},
        "format": "pdf",
    },
    ".docx": {
        "mimeTypes": {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/zip",
            "application/octet-stream",
        },
        "format": "docx",
    },
    ".txt": {
        "mimeTypes": {"text/plain", "application/octet-stream"},
        "format": "txt",
    },
}

SECTION_ALIASES = {
    "education": {"education", "academic background", "academics"},
    "experience": {
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "work history",
    },
    "projects": {"projects", "selected projects", "project experience"},
    "skills": {"skills", "technical skills", "core skills", "competencies"},
    "certifications": {"certifications", "certificates", "licenses"},
    "achievements": {"achievements", "awards", "honors", "accomplishments"},
    "interests": {"interests", "professional interests"},
}

DATE_PATTERN = re.compile(
    r"\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
    r"jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|"
    r"dec(?:ember)?)[\s.,-]+\d{4}\b|\b(?:19|20)\d{2}\b|\bpresent\b",
    re.IGNORECASE,
)
LOCATION_PATTERN = re.compile(
    r"\b[A-Z][A-Za-z .'-]+,\s*(?:[A-Z]{2}|United States|USA|Canada|"
    r"United Kingdom|UK|India|Australia|Germany|France|Singapore)\b"
)
ACTION_VERBS = (
    "built",
    "created",
    "developed",
    "designed",
    "managed",
    "led",
    "analyzed",
    "implemented",
    "coordinated",
    "improved",
    "launched",
    "researched",
    "supported",
    "automated",
    "increased",
    "reduced",
    "delivered",
)
INDUSTRY_KEYWORDS = {
    "Technology": ["technology", "software", "saas", "tech"],
    "Finance": ["finance", "financial", "bank", "investment", "fintech"],
    "Healthcare": ["healthcare", "health", "medical", "hospital"],
    "Retail": ["retail", "commerce", "ecommerce"],
    "Education": ["education", "school", "university", "teaching"],
    "Government": ["government", "public sector", "civic"],
    "Marketing": ["marketing", "advertising", "brand"],
}


class ResumeParser:
    def parse_upload(
        self,
        upload: FileStorage | None,
        *,
        source: str = "resume",
    ) -> dict[str, Any]:
        if upload is None:
            raise ApiError(
                "MISSING_FILE",
                "Choose a PDF, DOCX, or TXT file to continue.",
                details={"field": "file"},
            )

        supplied_name = upload.filename or ""
        extension = Path(supplied_name).suffix.lower()
        if extension not in SUPPORTED_TYPES:
            raise ApiError(
                "UNSUPPORTED_FILE_TYPE",
                "Only PDF, DOCX, and TXT files are supported.",
                details={"supportedExtensions": sorted(SUPPORTED_TYPES)},
            )

        declared_mime = (upload.mimetype or "").lower()
        if declared_mime not in SUPPORTED_TYPES[extension]["mimeTypes"]:
            raise ApiError(
                "UNSUPPORTED_FILE_TYPE",
                "The uploaded file type does not match a supported resume format.",
                details={
                    "declaredType": declared_mime,
                    "supportedExtensions": sorted(SUPPORTED_TYPES),
                },
            )

        data = upload.stream.read(MAX_UPLOAD_BYTES + 1)
        if not data:
            raise ApiError(
                "EMPTY_FILE",
                "The selected file is empty.",
                details={"maxBytes": MAX_UPLOAD_BYTES},
            )
        if len(data) > MAX_UPLOAD_BYTES:
            raise ApiError(
                "FILE_TOO_LARGE",
                "The selected file is larger than the 5 MB limit.",
                details={"maxBytes": MAX_UPLOAD_BYTES, "receivedBytes": len(data)},
            )

        file_format = str(SUPPORTED_TYPES[extension]["format"])
        self._validate_signature(file_format, data)
        try:
            text = self._extract_text(file_format, data)
        except ApiError:
            raise
        except Exception as error:
            raise ApiError(
                "UNREADABLE_FILE",
                "Path[IN] could not read this file. Try exporting it again or use TXT.",
                details={"format": file_format},
            ) from error

        text = self._clean_text(text)
        if len(text) < 20 or len(re.sub(r"\W", "", text)) < 10:
            raise ApiError(
                "NO_READABLE_TEXT",
                "No usable resume text was found in the file.",
                details={
                    "format": file_format,
                    "suggestion": "For scanned PDFs, export a text-based PDF or upload TXT.",
                },
            )

        fields = self._extract_fields(text, source=source)
        usable_count = sum(
            len(fields.get(category, []))
            for category in ("education", "roles", "responsibilities", "projects", "skills")
        )
        if usable_count == 0:
            raise ApiError(
                "NO_PROFILE_FACTS",
                "The file contained text, but no usable profile facts could be identified.",
                details={"format": file_format},
            )

        safe_stem = re.sub(r"[^A-Za-z0-9._ -]+", "", Path(supplied_name).stem)[:80]
        return {
            "file": {
                "displayName": f"{safe_stem or 'uploaded-profile'}{extension}",
                "format": file_format,
                "sizeBytes": len(data),
                "sha256": hashlib.sha256(data).hexdigest(),
                "retention": "Processed in memory and not permanently stored.",
            },
            "fields": fields,
            "summary": {
                "characterCount": len(text),
                "fieldCount": sum(len(items) for items in fields.values()),
                "explicitFactCount": sum(
                    1
                    for items in fields.values()
                    for item in items
                    if item["explicit"]
                ),
                "inferredSkillCount": sum(
                    1 for item in fields["skills"] if not item["explicit"]
                ),
            },
            "warnings": [
                "Review and correct every imported field before generating paths.",
                "Inferred skills are labeled and can be disabled or removed.",
            ],
        }

    @staticmethod
    def _validate_signature(file_format: str, data: bytes) -> None:
        if file_format == "pdf" and not data.startswith(b"%PDF-"):
            raise ApiError(
                "FILE_SIGNATURE_MISMATCH",
                "This file is not a valid PDF.",
                details={"format": file_format},
            )
        if file_format == "docx":
            if not data.startswith(b"PK"):
                raise ApiError(
                    "FILE_SIGNATURE_MISMATCH",
                    "This file is not a valid DOCX document.",
                    details={"format": file_format},
                )
            try:
                with zipfile.ZipFile(io.BytesIO(data)) as archive:
                    entries = archive.infolist()
                    names = {entry.filename for entry in entries}
            except zipfile.BadZipFile as error:
                raise ApiError(
                    "FILE_SIGNATURE_MISMATCH",
                    "This file is not a valid DOCX document.",
                    details={"format": file_format},
                ) from error
            if (
                len(entries) > MAX_DOCX_ENTRIES
                or sum(entry.file_size for entry in entries)
                > MAX_DOCX_UNCOMPRESSED_BYTES
            ):
                raise ApiError(
                    "UNSAFE_DOCUMENT",
                    "This DOCX expands beyond the safe document limit.",
                    details={
                        "format": file_format,
                        "maxEntries": MAX_DOCX_ENTRIES,
                        "maxUncompressedBytes": MAX_DOCX_UNCOMPRESSED_BYTES,
                    },
                )
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ApiError(
                    "FILE_SIGNATURE_MISMATCH",
                    "This archive is not a Word DOCX document.",
                    details={"format": file_format},
                )
        if file_format == "txt" and b"\x00" in data[:4096]:
            raise ApiError(
                "FILE_SIGNATURE_MISMATCH",
                "This TXT file appears to contain binary data.",
                details={"format": file_format},
            )

    @staticmethod
    def _extract_text(file_format: str, data: bytes) -> str:
        if file_format == "pdf":
            reader = PdfReader(io.BytesIO(data), strict=False)
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception as error:
                    raise ApiError(
                        "ENCRYPTED_FILE",
                        "Password-protected PDFs are not supported.",
                        details={"format": file_format},
                    ) from error
            if len(reader.pages) > MAX_PDF_PAGES:
                raise ApiError(
                    "UNSAFE_DOCUMENT",
                    "This PDF has more pages than Path[IN] can safely process.",
                    details={
                        "format": file_format,
                        "maxPages": MAX_PDF_PAGES,
                    },
                )
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if file_format == "docx":
            document = Document(io.BytesIO(data))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    paragraphs.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(paragraphs)
        try:
            return data.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise ApiError(
                "UNREADABLE_FILE",
                "TXT resumes must use UTF-8 text encoding.",
                details={"format": file_format},
            ) from error

    @staticmethod
    def _clean_text(text: str) -> str:
        cleaned_lines = []
        for raw_line in text.replace("\r", "\n").splitlines():
            line = re.sub(r"[\t ]+", " ", raw_line).strip()
            if line:
                cleaned_lines.append(line)
        return "\n".join(cleaned_lines)[:MAX_EXTRACTED_CHARACTERS]

    def _extract_fields(
        self, text: str, *, source: str
    ) -> dict[str, list[dict[str, Any]]]:
        lines = text.splitlines()
        section: str | None = None
        section_lines: dict[str, list[str]] = {
            "education": [],
            "experience": [],
            "projects": [],
            "skills": [],
            "certifications": [],
            "achievements": [],
            "interests": [],
        }
        unsectioned: list[str] = []

        for line in lines:
            heading = self._section_heading(line)
            if heading:
                section = heading
                continue
            if section:
                section_lines[section].append(line)
            else:
                unsectioned.append(line)

        fields: dict[str, list[dict[str, Any]]] = {
            "education": [],
            "roles": [],
            "responsibilities": [],
            "dates": [],
            "projects": [],
            "skills": [],
            "certifications": [],
            "industries": [],
            "achievements": [],
            "interests": [],
            "locations": [],
        }

        for value in section_lines["education"]:
            self._append(fields["education"], value, "education", source, 0.94)
        for value in section_lines["projects"]:
            self._append(fields["projects"], value, "projects", source, 0.93)
        for value in section_lines["certifications"]:
            self._append(fields["certifications"], value, "certifications", source, 0.95)
        for value in section_lines["achievements"]:
            self._append(fields["achievements"], value, "achievements", source, 0.93)
        for value in section_lines["interests"]:
            for part in self._split_list(value):
                self._append(fields["interests"], part, "interests", source, 0.9)

        experience_lines = section_lines["experience"]
        known_titles = [item["title"] for item in OCCUPATIONAL_TAXONOMY]
        for line in experience_lines:
            normalized_role = normalize_title(line)
            is_role_line = normalized_role or any(
                title.lower() in line.lower() for title in known_titles
            )
            if is_role_line:
                self._append(fields["roles"], line, "roles", source, 0.9)
            if is_role_line:
                continue
            if line.lower().lstrip("•- ").startswith(ACTION_VERBS):
                self._append(
                    fields["responsibilities"],
                    line.lstrip("•- "),
                    "responsibilities",
                    source,
                    0.92,
                )
            elif len(line.split()) >= 7:
                self._append(
                    fields["responsibilities"],
                    line.lstrip("•- "),
                    "responsibilities",
                    source,
                    0.78,
                )

        for line in section_lines["skills"]:
            for skill in self._split_list(line):
                self._append(fields["skills"], skill, "skills", source, 0.96)

        for match in DATE_PATTERN.findall(text):
            self._append(fields["dates"], match, "dates", source, 0.94)
        for match in LOCATION_PATTERN.findall(text):
            self._append(fields["locations"], match, "locations", source, 0.86)

        lower_text = "\n".join(
            [
                *unsectioned,
                *[
                    value
                    for values in section_lines.values()
                    for value in values
                ],
            ]
        ).lower()
        for industry, keywords in INDUSTRY_KEYWORDS.items():
            matching_keyword = next(
                (keyword for keyword in keywords if keyword in lower_text),
                None,
            )
            if matching_keyword:
                self._append(
                    fields["industries"],
                    industry,
                    "industries",
                    "inferred",
                    0.62,
                    explicit=False,
                    evidence=matching_keyword,
                    original_source=source,
                )

        if re.search(r"\bremote\b", lower_text):
            self._append(
                fields["locations"],
                "Remote",
                "locations",
                source,
                0.9,
            )

        explicit_skill_values = {
            item["normalized"] for item in fields["skills"]
        }
        for concept, aliases in CONCEPT_ALIASES.items():
            matching_alias = next((alias for alias in aliases if alias in lower_text), None)
            if not matching_alias:
                continue
            display = concept.title()
            normalized = self._normalize(display)
            if normalized in explicit_skill_values:
                continue
            explicit = any(
                matching_alias in line.lower() for line in section_lines["skills"]
            )
            self._append(
                fields["skills"],
                display,
                "skills",
                source if explicit else "inferred",
                0.88 if explicit else 0.58,
                explicit=explicit,
                evidence=matching_alias,
                original_source=source,
            )

        if not fields["roles"]:
            for line in unsectioned + experience_lines:
                if normalize_title(line):
                    self._append(fields["roles"], line, "roles", source, 0.72)
        return fields

    @staticmethod
    def _section_heading(line: str) -> str | None:
        normalized = re.sub(r"[^a-z ]", "", line.lower()).strip()
        if len(normalized.split()) > 4:
            return None
        for section, aliases in SECTION_ALIASES.items():
            if normalized in aliases:
                return section
        return None

    @staticmethod
    def _split_list(value: str) -> list[str]:
        return [
            item.strip(" •-\t")
            for item in re.split(r"[,;|•]", value)
            if item.strip(" •-\t")
        ]

    @staticmethod
    def _normalize(value: str) -> str:
        return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()

    def _append(
        self,
        target: list[dict[str, Any]],
        value: str,
        category: str,
        source: str,
        confidence: float,
        *,
        explicit: bool = True,
        evidence: str | None = None,
        original_source: str | None = None,
    ) -> None:
        cleaned = re.sub(r"\s+", " ", value).strip(" •-\t")
        if not cleaned or len(cleaned) > 500:
            return
        normalized = self._normalize(cleaned)
        if not normalized or any(item["normalized"] == normalized for item in target):
            return
        target.append(
            {
                "id": hashlib.sha1(
                    f"{source}:{category}:{normalized}".encode("utf-8")
                ).hexdigest()[:14],
                "value": cleaned,
                "normalized": normalized,
                "source": source,
                "confidence": round(confidence, 2),
                "explicit": explicit,
                "enabled": True,
                **({"evidence": evidence} if evidence else {}),
                **(
                    {"originalSource": original_source}
                    if original_source
                    else {}
                ),
            }
        )
